"""
Orion Stats - Word Export Service
Generates DOCX executive reports with Orion branding and full statistics.
"""
import io
from pathlib import Path
from typing import Optional, Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

from app.schemas.schemas import FilterCondition
from app.services.crosstab_service import calculate_crosstab
from app.services.stats_service import calculate_descriptive_stats


def _resolve_logo_path() -> Optional[Path]:
    """Resolve ORION wordmark path for report branding."""
    here = Path(__file__).resolve()
    candidates = [
        here.parents[1] / "assets" / "orion-wordmark-only.png",  # backend/app/assets
        here.parents[3] / "frontend" / "public" / "orion-wordmark-only.png",
        here.parents[3] / "Logo Orion Insights sem fundo.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _safe_num(value) -> Optional[float]:
    """Convert any numeric value to finite float."""
    if value is None:
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    if pd.isna(number) or number in (float("inf"), float("-inf")):
        return None
    return number


def _fmt(value, digits: int = 4) -> str:
    """Format numeric values for report tables."""
    number = _safe_num(value)
    if number is None:
        return "-"
    return f"{number:.{digits}f}"


def _format_filter_summary(filters: Optional[list[FilterCondition]], columns_meta: dict[str, str]) -> str:
    """Readable filter description for report header."""
    if not filters:
        return "Sem filtros ativos"
    chunks = []
    for condition in filters:
        col_name = columns_meta.get(condition.col_key, condition.col_key)
        chunks.append(f"{col_name} ({len(condition.values)} valor(es))")
    return "; ".join(chunks)


def _apply_base_style(document: Document):
    """Apply base typography style."""
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10)


def _add_logo_header(document: Document):
    """Insert Orion wordmark at top of document when available."""
    logo_path = _resolve_logo_path()
    if not logo_path:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    try:
        run.add_picture(str(logo_path), width=Inches(2.2))
    except Exception:
        # Keep export resilient if image loading fails.
        return


def _add_section_title(document: Document, title: str, level: int = 1):
    """Add standardized section titles."""
    heading = document.add_heading(title, level=level)
    heading_run = heading.runs[0] if heading.runs else heading.add_run(title)
    heading_run.font.color.rgb = RGBColor(27, 42, 74)
    heading_run.font.bold = True


def _add_table(document: Document, headers: list[str], rows: list[list[Any]]):
    """Add table with grid style."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        for run in header_cells[index].paragraphs[0].runs:
            run.font.bold = True

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = "" if value is None else str(value)
    return table


def _build_executive_rows(grouped, variables: list[str], columns_meta: dict[str, str]) -> list[list[Any]]:
    """Executive summary rows per selected variable."""
    rows = []
    if not grouped:
        return rows

    for variable in variables:
        variable_name = columns_meta.get(variable, variable)
        entries = []
        for group_key, group_stats in grouped.items():
            stat = next((item for item in group_stats if item.col_key == variable), None)
            if not stat or stat.mean is None:
                continue
            entries.append((group_key, stat))

        if not entries:
            continue

        best = max(entries, key=lambda item: item[1].mean if item[1].mean is not None else float("-inf"))
        worst = min(entries, key=lambda item: item[1].mean if item[1].mean is not None else float("inf"))
        highest_cv = max(entries, key=lambda item: item[1].cv if item[1].cv is not None else float("-inf"))
        highest_std = max(entries, key=lambda item: item[1].std if item[1].std is not None else float("-inf"))
        largest_n = max(entries, key=lambda item: item[1].count if item[1].count is not None else 0)
        spread = (best[1].mean - worst[1].mean) if best[1].mean is not None and worst[1].mean is not None else None

        rows.append([
            variable_name,
            best[0],
            _fmt(best[1].mean),
            worst[0],
            _fmt(worst[1].mean),
            _fmt(spread),
            highest_std[0],
            _fmt(highest_std[1].std),
            highest_cv[0],
            _fmt(highest_cv[1].cv, 2),
            largest_n[0],
            largest_n[1].count,
        ])

    return rows


def _build_group_summary_rows(group_summaries) -> list[list[Any]]:
    """Rows for grouped summary table."""
    rows = []
    for summary in group_summaries or []:
        rows.append([
            summary.group_key,
            summary.sample_size,
            _fmt(summary.pct_of_total, 2),
        ])
    return rows


def _build_group_detailed_rows(grouped, group_summaries) -> list[list[Any]]:
    """Rows for complete group-by-variable statistics table."""
    group_n_map = {}
    group_pct_map = {}
    for summary in group_summaries or []:
        group_n_map[summary.group_key] = summary.sample_size
        group_pct_map[summary.group_key] = summary.pct_of_total

    rows = []
    for group_key, group_stats in grouped.items():
        for stat in group_stats:
            rows.append([
                group_key,
                group_n_map.get(group_key, "-"),
                _fmt(group_pct_map.get(group_key), 2),
                stat.name,
                stat.count,
                stat.missing_count,
                _fmt(stat.missing_pct, 2),
                _fmt(stat.mean),
                _fmt(stat.median),
                _fmt(stat.std),
                _fmt(stat.cv, 2),
                _fmt(stat.min),
                _fmt(stat.max),
                _fmt(stat.range),
                _fmt(stat.q1),
                _fmt(stat.q3),
                _fmt(stat.iqr),
                _fmt(stat.p5),
                _fmt(stat.p95),
                _fmt(stat.skewness),
                _fmt(stat.kurtosis),
                _fmt(stat.ci_lower),
                _fmt(stat.ci_upper),
                _fmt(stat.sum),
            ])
    return rows


def _build_overall_rows(stats) -> list[list[Any]]:
    """Rows for overall descriptive statistics."""
    rows = []
    for stat in stats:
        rows.append([
            stat.name,
            stat.count,
            stat.missing_count,
            _fmt(stat.missing_pct, 2),
            _fmt(stat.mean),
            _fmt(stat.median),
            _fmt(stat.std),
            _fmt(stat.cv, 2),
            _fmt(stat.min),
            _fmt(stat.max),
            _fmt(stat.range),
            _fmt(stat.q1),
            _fmt(stat.q3),
            _fmt(stat.iqr),
            _fmt(stat.p5),
            _fmt(stat.p95),
            _fmt(stat.skewness),
            _fmt(stat.kurtosis),
            _fmt(stat.ci_lower),
            _fmt(stat.ci_upper),
            _fmt(stat.sum),
        ])
    return rows


def _build_test_rows(tests) -> list[list[Any]]:
    """Rows for group comparison tests."""
    rows = []
    for test in tests or []:
        rows.append([
            test.variable_name,
            test.test_name_display,
            _fmt(test.statistic),
            _fmt(test.p_value, 6),
            "Sim" if test.significant else "Nao",
            f"{test.effect_size_name} = {_fmt(test.effect_size)}" if test.effect_size is not None else "-",
            test.interpretation,
        ])
    return rows


def _to_filter_conditions(raw_filters: Optional[list[Any]]) -> list[FilterCondition]:
    """Normalize raw filter payload into schema-safe FilterCondition list."""
    normalized = []
    for item in raw_filters or []:
        if isinstance(item, FilterCondition):
            normalized.append(item)
            continue
        if not isinstance(item, dict):
            continue
        col_key = item.get("col_key")
        values = item.get("values")
        if not col_key or not isinstance(values, list):
            continue
        try:
            normalized.append(FilterCondition(col_key=str(col_key), values=values))
        except Exception:
            continue
    return normalized


def _safe_int(value: Any, default: int) -> int:
    try:
        parsed = int(value)
        return parsed if parsed > 0 else default
    except (TypeError, ValueError):
        return default


def _safe_float_conf(value: Any, default: float = 0.95) -> float:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return default
    if parsed <= 0 or parsed >= 1:
        return default
    return parsed


def _add_descriptive_composite_block(
    document: Document,
    df: pd.DataFrame,
    columns_meta: dict[str, str],
    payload: dict[str, Any],
    default_filters: Optional[list[FilterCondition]],
    default_treat_missing: bool,
):
    """Render one descriptive section in composite report."""
    variables = payload.get("variables") if isinstance(payload.get("variables"), list) else []
    variables = [str(variable) for variable in variables if str(variable) in df.columns]
    if not variables:
        document.add_paragraph("Bloco ignorado: nenhuma variavel valida encontrada para analise descritiva.")
        return

    group_by = payload.get("group_by") if isinstance(payload.get("group_by"), list) else []
    group_by = [str(group) for group in group_by if str(group) in df.columns]
    filters = _to_filter_conditions(payload.get("filters"))
    if not filters:
        filters = default_filters or []

    treat_missing_as_zero = payload.get("treat_missing_as_zero", default_treat_missing)
    run_comparison_tests = bool(payload.get("run_comparison_tests", True if group_by else False))
    confidence_level = _safe_float_conf(payload.get("confidence_level"), default=0.95)
    max_groups = _safe_int(payload.get("max_groups"), default=200)

    sample_size, stats, grouped, summaries, tests, total_groups = calculate_descriptive_stats(
        df=df,
        variables=variables,
        columns_meta=columns_meta,
        filters=filters if filters else None,
        group_by=group_by if group_by else None,
        treat_missing_as_zero=bool(treat_missing_as_zero),
        confidence_level=confidence_level,
        run_comparison_tests=run_comparison_tests,
        max_groups=max_groups,
    )

    scope_rows = [
        ["Amostra", sample_size],
        ["Variaveis", ", ".join(columns_meta.get(variable, variable) for variable in variables)],
        ["Agrupamento", ", ".join(columns_meta.get(group, group) for group in group_by) if group_by else "Sem agrupamento"],
        ["Total de grupos", total_groups if total_groups is not None else "-"],
        ["Filtros", _format_filter_summary(filters, columns_meta)],
    ]
    _add_table(document, ["Indicador", "Valor"], scope_rows)

    overall_headers = [
        "Variavel", "N", "Ausentes", "% Ausentes", "Media", "Mediana", "DP", "CV%",
        "Min", "Max", "Amplitude", "Q1", "Q3", "IQR", "P5", "P95",
    ]
    overall_rows = []
    for stat in stats:
        overall_rows.append([
            stat.name,
            stat.count,
            stat.missing_count,
            _fmt(stat.missing_pct, 2),
            _fmt(stat.mean),
            _fmt(stat.median),
            _fmt(stat.std),
            _fmt(stat.cv, 2),
            _fmt(stat.min),
            _fmt(stat.max),
            _fmt(stat.range),
            _fmt(stat.q1),
            _fmt(stat.q3),
            _fmt(stat.iqr),
            _fmt(stat.p5),
            _fmt(stat.p95),
        ])
    _add_table(document, overall_headers, overall_rows if overall_rows else [["-"] * len(overall_headers)])

    if grouped:
        summary_headers = ["Grupo", "N Grupo", "% Total"]
        summary_rows = _build_group_summary_rows(summaries)
        _add_table(document, summary_headers, summary_rows if summary_rows else [["-", "-", "-"]])

    if tests:
        test_headers = ["Variavel", "Teste", "Estatistica", "p-valor", "Significativo", "Tamanho de efeito"]
        test_rows = []
        for test in tests:
            test_rows.append([
                test.variable_name,
                test.test_name_display,
                _fmt(test.statistic),
                _fmt(test.p_value, 6),
                "Sim" if test.significant else "Nao",
                f"{test.effect_size_name} = {_fmt(test.effect_size)}" if test.effect_size is not None else "-",
            ])
        _add_table(document, test_headers, test_rows if test_rows else [["-"] * len(test_headers)])


def _add_crosstab_composite_block(
    document: Document,
    df: pd.DataFrame,
    columns_meta: dict[str, str],
    payload: dict[str, Any],
    default_filters: Optional[list[FilterCondition]],
    default_treat_missing: bool,
):
    """Render one crosstab section in composite report."""
    row_variable = str(payload.get("row_variable", "")).strip()
    col_variable = str(payload.get("col_variable", "")).strip()
    if not row_variable or not col_variable:
        document.add_paragraph("Bloco ignorado: row_variable e col_variable sao obrigatorias no crosstab.")
        return
    if row_variable not in df.columns or col_variable not in df.columns:
        document.add_paragraph("Bloco ignorado: variaveis do crosstab nao encontradas no dataset atual.")
        return

    filters = _to_filter_conditions(payload.get("filters"))
    if not filters:
        filters = default_filters or []
    treat_missing_as_zero = payload.get("treat_missing_as_zero", default_treat_missing)
    max_rows = _safe_int(payload.get("max_rows"), default=30)
    max_cols = _safe_int(payload.get("max_cols"), default=30)

    result = calculate_crosstab(
        df=df,
        row_variable=row_variable,
        col_variable=col_variable,
        columns_meta=columns_meta,
        filters=filters if filters else None,
        max_rows=max_rows,
        max_cols=max_cols,
        treat_missing_as_zero=bool(treat_missing_as_zero),
    )

    info_rows = [
        ["Variavel de linhas", result.row_variable_name],
        ["Variavel de colunas", result.col_variable_name],
        ["Amostra", result.sample_size],
        ["Total da tabela", result.grand_total],
        ["Filtros", _format_filter_summary(filters, columns_meta)],
    ]
    _add_table(document, ["Indicador", "Valor"], info_rows)

    headers = [result.row_variable_name] + result.col_labels + ["Total Linha"]
    table_rows = []
    for row_idx, row_label in enumerate(result.row_labels):
        values = [f"{count} ({result.percentages[row_idx][col_idx]}%)" for col_idx, count in enumerate(result.counts[row_idx])]
        table_rows.append([row_label, *values, result.row_totals[row_idx]])
    totals_row = ["Total Coluna", *result.col_totals, result.grand_total]
    table_rows.append(totals_row)
    _add_table(document, headers, table_rows)

    if result.chi_square is not None:
        test_rows = [
            ["X²", _fmt(result.chi_square)],
            ["p-valor", _fmt(result.chi_square_p_value, 6)],
            ["V de Cramer", _fmt(result.cramers_v, 4)],
            ["Interpretacao", result.interpretation or "-"],
        ]
        _add_table(document, ["Teste Qui-Quadrado", "Resultado"], test_rows)


def _add_ml_composite_block(
    document: Document,
    payload: dict[str, Any],
):
    """Render one ML section from payload snapshot."""
    snapshot = payload.get("result_snapshot") if isinstance(payload.get("result_snapshot"), dict) else {}
    target_name = payload.get("target_name") or payload.get("target") or "-"
    feature_names = payload.get("feature_names") if isinstance(payload.get("feature_names"), list) else payload.get("features", [])
    features_str = ", ".join(str(item) for item in feature_names) if feature_names else "-"

    info_rows = [
        ["Variavel alvo", target_name],
        ["Features", features_str],
        ["Metrica de selecao", str(payload.get("selection_metric", "-")).upper()],
        ["Modelo selecionado", payload.get("selected_model") or snapshot.get("best_model_label") or "-"],
        ["Model ID", snapshot.get("model_id") or "-"],
    ]
    _add_table(document, ["Configuracao ML", "Valor"], info_rows)

    models = snapshot.get("models") if isinstance(snapshot.get("models"), list) else []
    if models:
        metric_headers = ["Modelo", "R2", "RMSE", "MAE", "MAPE", "Melhor"]
        metric_rows = []
        for model in models:
            if not isinstance(model, dict):
                continue
            metric_rows.append([
                model.get("label", "-"),
                _fmt(model.get("r2")),
                _fmt(model.get("rmse")),
                _fmt(model.get("mae")),
                _fmt(model.get("mape"), 2),
                "Sim" if bool(model.get("is_best")) else "Nao",
            ])
        _add_table(document, metric_headers, metric_rows if metric_rows else [["-"] * len(metric_headers)])

    linear = snapshot.get("linear_regression") if isinstance(snapshot.get("linear_regression"), dict) else {}
    if linear:
        linear_rows = [
            ["Equacao", linear.get("equation", "-")],
            ["R2", _fmt(linear.get("r2"))],
            ["RMSE", _fmt(linear.get("rmse"))],
            ["Intercepto", _fmt(linear.get("intercept"))],
        ]
        _add_table(document, ["Regressao Linear", "Valor"], linear_rows)

        coefficients = linear.get("coefficients") if isinstance(linear.get("coefficients"), list) else []
        if coefficients:
            coef_headers = ["Feature", "Coef.", "Erro Padrao", "t-valor", "p-valor"]
            coef_rows = []
            for coef in coefficients[:30]:
                if not isinstance(coef, dict):
                    continue
                coef_rows.append([
                    coef.get("feature", "-"),
                    _fmt(coef.get("coefficient")),
                    _fmt(coef.get("std_error")),
                    _fmt(coef.get("t_value")),
                    _fmt(coef.get("p_value"), 6),
                ])
            _add_table(document, coef_headers, coef_rows if coef_rows else [["-"] * len(coef_headers)])


def _add_composite_sections(
    document: Document,
    df: pd.DataFrame,
    columns_meta: dict[str, str],
    report_sections: Optional[list[dict[str, Any]]],
    default_filters: Optional[list[FilterCondition]],
    default_treat_missing: bool,
):
    """Render user-composed analysis blocks inside the exported report."""
    if not report_sections:
        return

    _add_section_title(document, "8. Blocos Compostos do Relatorio", level=1)

    valid_sections = [section for section in report_sections if isinstance(section, dict)]
    if not valid_sections:
        document.add_paragraph("Nenhum bloco valido encontrado no relatorio composto.")
        return

    for index, section in enumerate(valid_sections, start=1):
        section_type = str(section.get("section_type", "")).strip().lower()
        title = str(section.get("title", f"Bloco {index}"))
        created_at = section.get("created_at")
        payload = section.get("payload") if isinstance(section.get("payload"), dict) else {}

        _add_section_title(document, f"8.{index} {title}", level=2)
        if created_at:
            document.add_paragraph(f"Criado em: {created_at}")
        document.add_paragraph(f"Tipo: {section_type or '-'}")

        try:
            if section_type == "descriptive":
                _add_descriptive_composite_block(
                    document=document,
                    df=df,
                    columns_meta=columns_meta,
                    payload=payload,
                    default_filters=default_filters,
                    default_treat_missing=default_treat_missing,
                )
            elif section_type == "crosstab":
                _add_crosstab_composite_block(
                    document=document,
                    df=df,
                    columns_meta=columns_meta,
                    payload=payload,
                    default_filters=default_filters,
                    default_treat_missing=default_treat_missing,
                )
            elif section_type == "ml":
                _add_ml_composite_block(
                    document=document,
                    payload=payload,
                )
            else:
                document.add_paragraph("Tipo de bloco nao suportado neste export.")
        except Exception as error:
            document.add_paragraph(f"Falha ao renderizar bloco '{title}': {str(error)}")


def create_word_export(
    df: pd.DataFrame,
    variables: list[str],
    columns_meta: dict[str, str],
    filters: Optional[list[FilterCondition]] = None,
    group_by: Optional[list[str]] = None,
    treat_missing_as_zero: bool = True,
    report_sections: Optional[list[dict[str, Any]]] = None,
) -> io.BytesIO:
    """Create DOCX report with complete statistics and executive interpretation."""
    sample_size, stats, grouped, summaries, tests, total_groups = calculate_descriptive_stats(
        df=df,
        variables=variables,
        columns_meta=columns_meta,
        filters=filters,
        group_by=group_by,
        treat_missing_as_zero=treat_missing_as_zero,
        run_comparison_tests=True if group_by else False,
        max_groups=200,
    )

    document = Document()
    _apply_base_style(document)
    _add_logo_header(document)

    title = document.add_heading("Relatorio Executivo Orion - Estatisticas", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.runs[0] if title.runs else title.add_run("Relatorio Executivo Orion - Estatisticas")
    title_run.font.color.rgb = RGBColor(27, 42, 74)

    subtitle = document.add_paragraph(f"Gerado em {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M:%S')}")
    subtitle.runs[0].font.size = Pt(10)

    _add_section_title(document, "1. Escopo da Analise", level=1)
    scope_headers = ["Indicador", "Valor"]
    scope_rows = [
        ["Amostra total", sample_size],
        ["Variaveis analisadas", ", ".join(columns_meta.get(variable, variable) for variable in variables)],
        ["Agrupamento", ", ".join(columns_meta.get(group, group) for group in (group_by or [])) if group_by else "Sem agrupamento"],
        ["Total de grupos", total_groups if total_groups is not None else "-"],
        ["Filtros ativos", _format_filter_summary(filters, columns_meta)],
        ["Tratamento de ausentes", "Ausentes tratados como 0" if treat_missing_as_zero else "Ausentes removidos por variavel"],
    ]
    _add_table(document, scope_headers, scope_rows)

    if grouped:
        _add_section_title(document, "2. Resumo Executivo por Variavel", level=1)
        executive_headers = [
            "Variavel",
            "Melhor Grupo",
            "Media Melhor",
            "Pior Grupo",
            "Media Pior",
            "Amplitude",
            "Maior DP Grupo",
            "DP",
            "Maior CV Grupo",
            "CV%",
            "Maior N Grupo",
            "N",
        ]
        executive_rows = _build_executive_rows(grouped, variables, columns_meta)
        _add_table(document, executive_headers, executive_rows if executive_rows else [["-", "-", "-", "-", "-", "-", "-", "-", "-", "-", "-", "-"]])

    _add_section_title(document, "3. Estatisticas Gerais (todas as variaveis)", level=1)
    overall_headers = [
        "Variavel", "N", "Ausentes", "% Ausentes", "Media", "Mediana", "DP", "CV%",
        "Min", "Max", "Amplitude", "Q1", "Q3", "IQR", "P5", "P95",
        "Assimetria", "Curtose", "IC Inf", "IC Sup", "Soma"
    ]
    overall_rows = _build_overall_rows(stats)
    _add_table(document, overall_headers, overall_rows if overall_rows else [["-"] * len(overall_headers)])

    if grouped:
        _add_section_title(document, "4. Estatisticas por Grupo - Resumo", level=1)
        summary_headers = ["Grupo", "N Grupo", "% Total"]
        summary_rows = _build_group_summary_rows(summaries)
        _add_table(document, summary_headers, summary_rows if summary_rows else [["-", "-", "-"]])

        _add_section_title(document, "5. Estatisticas por Grupo - Detalhado", level=1)
        detailed_headers = [
            "Grupo", "N Grupo", "% Total", "Variavel", "N", "Ausentes", "% Ausentes",
            "Media", "Mediana", "DP", "CV%", "Min", "Max", "Amplitude",
            "Q1", "Q3", "IQR", "P5", "P95", "Assimetria", "Curtose", "IC Inf", "IC Sup", "Soma"
        ]
        detailed_rows = _build_group_detailed_rows(grouped, summaries)
        _add_table(document, detailed_headers, detailed_rows if detailed_rows else [["-"] * len(detailed_headers)])

    if tests:
        _add_section_title(document, "6. Testes Comparativos entre Grupos", level=1)
        test_headers = ["Variavel", "Teste", "Estatistica", "p-valor", "Significativo", "Tamanho de efeito", "Interpretacao"]
        test_rows = _build_test_rows(tests)
        _add_table(document, test_headers, test_rows if test_rows else [["-"] * len(test_headers)])

    _add_section_title(document, "7. Leitura Executiva", level=1)
    bullets = [
        "Compare primeiro Media, CV% e N para decidir onde agir.",
        "Diferenca estatistica (p-valor) deve ser avaliada junto com tamanho de efeito.",
        "Grupos com CV% alto indicam instabilidade operacional.",
        "Use o grupo com maior media e baixa variabilidade como benchmark.",
    ]
    for bullet in bullets:
        document.add_paragraph(bullet, style="List Bullet")

    _add_composite_sections(
        document=document,
        df=df,
        columns_meta=columns_meta,
        report_sections=report_sections,
        default_filters=filters,
        default_treat_missing=treat_missing_as_zero,
    )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
